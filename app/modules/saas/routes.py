from . import saas_bp
from flask import request, jsonify, render_template, send_from_directory, current_app, redirect, url_for, flash, session
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
import os
import subprocess
import shutil
import platform
from app.models import Conversion, Payment, Subscription
from app.payments.services import get_plan_rows, get_plan, MERCHANT_PHONE
from app.security import validate_csrf, premium_required
from app.utils.notifications import notify_admin_new_subscription


def _t(fr_text, en_text):
    return en_text if session.get('lang') == 'en' else fr_text

# Optional imports with fallback
try:
    from pdf2docx import Converter as PDF2DOCX_Converter
except ImportError:
    PDF2DOCX_Converter = None

try:
    from PyPDF2 import PdfReader, PdfWriter, PdfMerger
except ImportError:
    PdfReader = PdfWriter = PdfMerger = None

try:
    import qrcode
except ImportError:
    qrcode = None

import io
import base64
from datetime import datetime
import json

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml import OxmlElement
except ImportError:
    Document = None

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = ImageDraw = ImageFont = None

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)


import threading
import time as _time


def _libreoffice_convert(input_path, out_dir, to_format):
    """Convert a file using LibreOffice headless.
    
    On Windows, LibreOffice sometimes converts successfully but the process
    hangs instead of exiting. We work around this by:
    1. Launching the process
    2. Polling for the output file to appear
    3. Killing the process once the file is ready (or on timeout)
    """
    soffice = _find_soffice()
    if not soffice:
        return False
    # Resolve to long paths (avoid 8.3 short names like MONSIE~1)
    try:
        input_path = os.path.realpath(input_path)
        out_dir = os.path.realpath(out_dir)
    except Exception:
        pass

    # Determine expected output filename
    base = os.path.splitext(os.path.basename(input_path))[0]
    expected_out = os.path.join(out_dir, f"{base}.{to_format}")

    cmd = [soffice, '--headless', '--norestore', '--nofirststartwizard',
           '--convert-to', to_format, '--outdir', out_dir, input_path]
    try:
        proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

        # Poll for the output file with a total timeout of 90s
        max_wait = 90
        poll_interval = 1.0
        waited = 0.0
        while waited < max_wait:
            _time.sleep(poll_interval)
            waited += poll_interval
            # Check if the process ended on its own
            if proc.poll() is not None:
                break
            # Check if the output file appeared and has content
            if os.path.exists(expected_out) and os.path.getsize(expected_out) > 0:
                # Give LibreOffice a moment to finish writing
                _time.sleep(2)
                break

        # Read any output for logging
        try:
            stdout, stderr = proc.communicate(timeout=3)
        except Exception:
            stdout, stderr = b'', b''

        # Kill the process if it's still running (common on Windows)
        if proc.poll() is None:
            try:
                proc.kill()
                proc.wait(timeout=5)
            except Exception:
                pass

        success = os.path.exists(expected_out) and os.path.getsize(expected_out) > 0
        current_app.logger.info(
            f'LibreOffice convert: success={success} waited={waited:.0f}s '
            f'file={expected_out} '
            f'stdout={stdout.decode("utf-8", errors="replace")[:150]} '
            f'stderr={stderr.decode("utf-8", errors="replace")[:150]}'
        )
        return success
    except Exception as exc:
        current_app.logger.warning(f'LibreOffice convert error: {exc}')
        return False


def _find_soffice():
    # Check app config first
    try:
        cfg_path = current_app.config.get('LIBREOFFICE_PATH')
    except Exception:
        cfg_path = None
    candidates = []
    if cfg_path:
        candidates.append(cfg_path)
        # On Windows, also try .com version (console mode, better headless support)
        if cfg_path.endswith('.exe'):
            candidates.append(cfg_path.replace('.exe', '.com'))
    # Check system PATH
    which_soffice = shutil.which('soffice') or shutil.which('soffice.com') or shutil.which('soffice.bin')
    if which_soffice:
        candidates.append(which_soffice)
    # On Windows, common install locations — prefer .com (console) over .exe
    if platform.system().lower().startswith('win'):
        prog = os.environ.get('PROGRAMFILES', r"C:\Program Files")
        prog_x86 = os.environ.get('PROGRAMFILES(X86)', r"C:\Program Files (x86)")
        for base in [prog, prog_x86]:
            candidates.append(os.path.join(base, 'LibreOffice', 'program', 'soffice.com'))
            candidates.append(os.path.join(base, 'LibreOffice', 'program', 'soffice.exe'))
    # Validate candidates
    for c in candidates:
        if not c:
            continue
        try:
            if os.path.exists(c):
                return c
        except Exception:
            continue
    return None


def _generate_code_content(language: str, name: str, description: str):
        if language == 'html':
                content = f"""<!-- {description} -->
<!doctype html>
<html>
    <head>
        <meta charset=\"utf-8\"> 
        <title>{name}</title>
    </head>
    <body>
        <h1>{name}</h1>
        <p>{description}</p>
    </body>
</html>
"""
                filename = f"{name}.html"
        elif language == 'javascript' or language == 'js':
                content = f"""// {description}
function main() {{
    console.log('{description}');
}}

main();
"""
                filename = f"{name}.js"
        else:
                # default to python
                content = f"""# {description}
def main():
        print('{description}')

if __name__ == '__main__':
        main()
"""
                filename = f"{name}.py"
        return filename, content


@saas_bp.route('/convert', methods=['POST'])
@premium_required
def convert_file():
    f = request.files.get('file')
    direction = request.form.get('direction') or request.values.get('direction') or 'pdf-to-docx'
    if not f:
        return jsonify({'error': 'no file provided'}), 400
    filename = secure_filename(f.filename)
    path = os.path.join(UPLOAD_DIR, filename)
    f.save(path)

    name, ext = os.path.splitext(filename)
    ext = ext.lower()
    out_filename = None
    ok = False

    # Determine desired conversion
    if direction == 'pdf-to-docx':
        out_filename = f"{name}.docx"
        out_path = os.path.join(UPLOAD_DIR, out_filename)
        # Try libreoffice
        ok = _libreoffice_convert(path, UPLOAD_DIR, 'docx')
        if not ok:
            # Fallback to pdf2docx if installed
            try:
                cv = PDF2DOCX_Converter(path)
                cv.convert(out_path, start=0, end=None)
                cv.close()
                ok = os.path.exists(out_path)
            except Exception:
                ok = False
            if not ok:
                return jsonify({'error': 'LibreOffice unavailable and pdf2docx failed; pdf -> docx conversion not supported on server'}), 503
    elif direction == 'docx-to-pdf' or (ext in ['.doc', '.docx'] and direction == 'docx-to-pdf'):
        out_filename = f"{name}.pdf"
        out_path = os.path.join(UPLOAD_DIR, out_filename)
        # Try libreoffice for docx->pdf
        ok = _libreoffice_convert(path, UPLOAD_DIR, 'pdf')
        if not ok:
            # As fallback, try pypandoc if available
            try:
                import pypandoc
                pypandoc.convert_file(path, 'pdf', outputfile=out_path)
                ok = True
            except Exception:
                ok = False
        if not ok:
            return jsonify({'error': 'Conversion to PDF failed; install LibreOffice or pandoc on server'}), 503
    else:
        return jsonify({'error': 'unsupported conversion direction'}), 400

    # If libreoffice used, it may have created a file with same base name but different extension
    if os.path.exists(out_path):
        Conversion.create(original=filename, output=os.path.basename(out_path), direction=direction, status='done')
        return jsonify({'converted': True, 'original': filename, 'output': os.path.basename(out_path), 'download': f"/modules/saas/download/{os.path.basename(out_path)}"})
    else:
        return jsonify({'error': 'conversion failed'}), 500


@saas_bp.route('/pdf/merge', methods=['POST'])
@premium_required
def pdf_merge():
    if PdfMerger is None:
        return jsonify({'error': 'PDF merge not available on server'}), 503
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'no files provided'}), 400
    merger = PdfMerger()
    names = []
    for f in files:
        filename = secure_filename(f.filename)
        path = os.path.join(UPLOAD_DIR, filename)
        f.save(path)
        try:
            merger.append(path)
            names.append(filename)
        except Exception:
            continue
    out_name = f"merged_{int(datetime.utcnow().timestamp()*1000)}.pdf"
    out_path = os.path.join(UPLOAD_DIR, out_name)
    try:
        merger.write(out_path)
        merger.close()
        Conversion.create(original=','.join(names), output=out_name, direction='pdf-merge', status='done')
        return jsonify({'merged': True, 'output': out_name, 'download': f"/modules/saas/download/{out_name}"})
    except Exception:
        current_app.logger.exception('PDF merge failed')
        return jsonify({'error': 'merge failed'}), 500


@saas_bp.route('/pdf/split', methods=['POST'])
@premium_required
def pdf_split():
    if PdfReader is None or PdfWriter is None:
        return jsonify({'error': 'PDF split not available on server'}), 503
    f = request.files.get('file')
    if not f:
        return jsonify({'error': 'no file provided'}), 400
    filename = secure_filename(f.filename)
    path = os.path.join(UPLOAD_DIR, filename)
    f.save(path)
    try:
        reader = PdfReader(path)
        pages = []
        for i, page in enumerate(reader.pages, start=1):
            writer = PdfWriter()
            writer.add_page(page)
            out_name = f"{os.path.splitext(filename)[0]}_page_{i}.pdf"
            out_path = os.path.join(UPLOAD_DIR, out_name)
            with open(out_path, 'wb') as fh:
                writer.write(fh)
            pages.append({
                'output': out_name,
                'download': f"/modules/saas/download/{out_name}",
                'page': i
            })
        Conversion.create(original=filename, output=','.join([p['output'] for p in pages]), direction='pdf-split', status='done')
        return jsonify({'split': True, 'pages': pages})
    except Exception:
        current_app.logger.exception('PDF split failed')
        return jsonify({'error': 'split failed'}), 500


@saas_bp.route('/')
def index():
    active_subscription = None
    latest_payment = None
    if current_user.is_authenticated:
        active_subscription = Subscription.get_active_for_user(current_user.id)
        payments = Payment.all_for_user(current_user.id, limit=1)
        latest_payment = payments[0] if payments else None
    return render_template(
        'saas.html',
        plans=get_plan_rows(),
        merchant_phone=MERCHANT_PHONE,
        active_subscription=active_subscription,
        latest_payment=latest_payment
    )


@saas_bp.route('/pricing')
def pricing():
    return redirect(url_for('saas.index'))


@saas_bp.route('/subscribe/manual', methods=['POST'])
@login_required
def subscribe_manual():
    if not validate_csrf():
        flash(_t('Session expirée. Veuillez réessayer.', 'Session expired. Please try again.'))
        return redirect(url_for('saas.index'))

    plan_code = request.form.get('plan', '').strip()
    provider = request.form.get('provider', '').strip().lower()
    transaction_ref = request.form.get('transaction_ref', '').strip()
    payer_phone = request.form.get('payer_phone', '').strip()
    payer_name = request.form.get('payer_name', '').strip()

    if provider not in ['orange_money', 'wave']:
        flash(_t('Moyen de paiement invalide. Choisissez Orange Money ou Wave.', 'Invalid payment method. Choose Orange Money or Wave.'))
        return redirect(url_for('saas.index'))

    if not transaction_ref or not payer_phone:
        flash(_t('Référence de transaction et numéro payeur obligatoires.', 'Transaction reference and payer phone are required.'))
        return redirect(url_for('saas.index'))

    plan = get_plan(plan_code)
    payment = Payment.create_manual(
        user_id=current_user.id,
        plan=plan.code,
        provider=provider,
        amount=plan.amount_usd,
        transaction_ref=transaction_ref,
        payer_phone=payer_phone,
        payer_name=payer_name
    )
    Subscription.create(
        user_id=current_user.id,
        plan=plan.code,
        amount=plan.amount_usd,
        provider=provider,
        active=False,
        status='pending_validation',
        payment_id=payment['id'],
        duration_days=30,
        features=plan.highlights
    )

    # Alert admin on WhatsApp that a payment needs validation
    try:
        notify_admin_new_subscription(
            user_name=current_user.username,
            plan_code=plan.code,
            amount=plan.amount_usd,
            provider=provider,
            transaction_ref=transaction_ref,
            payer_phone=payer_phone
        )
    except Exception as exc:
        current_app.logger.warning(f'Admin WhatsApp notification failed: {exc}')

    flash(
        _t(
            f"Paiement soumis avec succès. Envoyez l'équivalent de ${plan.amount_usd} USD au {MERCHANT_PHONE} puis attendez la validation admin.",
            f"Payment submitted successfully. Send the equivalent of ${plan.amount_usd} USD to {MERCHANT_PHONE} and wait for admin approval."
        )
    )
    return redirect(url_for('saas.index'))


@saas_bp.route('/history')
@premium_required
def history():
    rows = Conversion.all(limit=50)
    return jsonify({'history': rows})


@saas_bp.route('/download/<path:filename>')
@premium_required
def download(filename):
    return send_from_directory(UPLOAD_DIR, filename, as_attachment=True)


@saas_bp.route('/cv/generate', methods=['POST'])
@premium_required
def generate_cv():
    profile = request.json or {}
    # Very simple CV generation from template
    cv_text = f"{profile.get('name','Nom')}\n{profile.get('title','Titre')}\n\n{profile.get('summary','Résumé')}"
    return jsonify({'cv': cv_text})


@saas_bp.route('/code/generate', methods=['POST'])
@premium_required
def generate_code():
    data = request.json or {}
    language = (data.get('language') or 'python').lower()
    name = secure_filename(data.get('name') or 'generated')
    description = data.get('description') or 'Generated code'
    filename, content = _generate_code_content(language, name, description)
    out_path = os.path.join(UPLOAD_DIR, filename)
    try:
        with open(out_path, 'w', encoding='utf-8') as fh:
            fh.write(content)
        Conversion.create(original=name, output=filename, direction='code-generate', status='done')
        return jsonify({'generated': True, 'output': filename, 'download': f"/modules/saas/download/{filename}"})
    except Exception:
        current_app.logger.exception('Code generation failed')
        return jsonify({'error': 'generation failed'}), 500


@saas_bp.route('/qr/generate', methods=['POST'])
@premium_required
def generate_qr():
    # Accept JSON or form data
    data = request.json or request.form or {}
    # If multiple fields provided, combine into text
    text = data.get('text')
    if not text:
        # assemble from fields
        parts = []
        for k, v in (data.items() if isinstance(data, dict) else []):
            parts.append(f"{k}:{v}")
        text = '\n'.join(parts) if parts else None
    if not text:
        return jsonify({'error': 'no data provided for QR code'}), 400
    try:
        img = qrcode.make(text)
        bio = io.BytesIO()
        img.save(bio, format='PNG')
        bio.seek(0)
        fname = f"qr_{int(datetime.utcnow().timestamp()*1000)}.png"
        out_path = os.path.join(UPLOAD_DIR, fname)
        with open(out_path, 'wb') as fh:
            fh.write(bio.getvalue())
        b64 = base64.b64encode(bio.getvalue()).decode('ascii')
        Conversion.create(original='qr', output=fname, direction='qr-generate', status='done')
        return jsonify({'generated': True, 'output': fname, 'download': f"/modules/saas/download/{fname}", 'data_uri': f"data:image/png;base64,{b64}"})
    except Exception:
        current_app.logger.exception('QR generation failed')
        return jsonify({'error': 'qr generation failed'}), 500


@saas_bp.route('/cv/professional', methods=['POST'])
@premium_required
def generate_professional_cv():
    # Accept JSON or form-data (with optional uploaded file 'photo')
    form = request.form or {}
    j = request.get_json(silent=True) or {}
    # prefer form values when present (FormData from client)
    def get_field(key, default=''):
        return (form.get(key) or j.get(key) or default)

    name = get_field('name', 'Nom')
    title = get_field('title', '')
    contact = get_field('contact', '')
    summary = get_field('summary', '')
    # experiences/education may be JSON strings in form
    def parse_list_field(val):
        if not val:
            return []
        if isinstance(val, list):
            return val
        try:
            return json.loads(val)
        except Exception:
            # fallback: comma-separated
            return [v.strip() for v in str(val).split(',') if v.strip()]

    experiences = parse_list_field(get_field('experiences', []))
    education = parse_list_field(get_field('education', []))
    skills = parse_list_field(get_field('skills', []))
    # handle uploaded photo
    photo = request.files.get('photo') if hasattr(request, 'files') else None
    photo_path = None
    if photo:
        try:
            photo_name = secure_filename(photo.filename)
            if photo_name:
                unique = f"{int(datetime.utcnow().timestamp()*1000)}_{photo_name}"
                photo_path = os.path.join(UPLOAD_DIR, unique)
                photo.save(photo_path)
        except Exception:
            photo_path = None
    try:
        doc = Document()
        # set document defaults
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # page margins
        section = doc.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

        # add page border via XML
        try:
            sectPr = section._sectPr
            pgBorders = OxmlElement('w:pgBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                node = OxmlElement(f'w:{side}')
                node.set('w:val', 'single')
                node.set('w:sz', '12')
                node.set('w:space', '24')
                node.set('w:color', 'B0CFFF')
                pgBorders.append(node)
            sectPr.append(pgBorders)
        except Exception:
            pass

        # Header: two-column table for photo + info
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(1.6)
        table.columns[1].width = Inches(4.8)

        cell_photo = table.cell(0, 0)
        cell_info = table.cell(0, 1)

        # Photo: use uploaded photo if provided, otherwise placeholder
        if photo_path and os.path.exists(photo_path):
            try:
                paragraph = cell_photo.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(photo_path, width=Inches(1.4))
            except Exception:
                cell_photo.text = '[Photo]'
        else:
            try:
                img = Image.new('RGB', (400,400), color=(230,230,230))
                draw = ImageDraw.Draw(img)
                txt = 'Photo'
                # try to load a default font
                try:
                    fnt = ImageFont.truetype('arial.ttf', 40)
                except Exception:
                    fnt = ImageFont.load_default()
                # Use textbbox (Pillow 10+) instead of removed textsize
                try:
                    bbox = draw.textbbox((0, 0), txt, font=fnt)
                    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
                except AttributeError:
                    w, h = 80, 40  # fallback estimate
                draw.rectangle([10,10,390,390], outline=(180,180,180), width=6)
                draw.text(((400-w)/2,(400-h)/2), txt, font=fnt, fill=(100,100,100))
                bio = io.BytesIO()
                img.save(bio, format='PNG')
                bio.seek(0)
                paragraph = cell_photo.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(bio, width=Inches(1.4))
            except Exception:
                cell_photo.text = '[Photo]'

        # Info: name, title, contact with emoji
        p = cell_info.paragraphs[0]
        name_run = p.add_run(name + ' ')
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.color.rgb = RGBColor(4,32,51)

        if title:
            p = cell_info.add_paragraph()
            tr = p.add_run(title)
            tr.italic = True
            tr.font.size = Pt(12)
            tr.font.color.rgb = RGBColor(80,80,80)

        if contact:
            p = cell_info.add_paragraph()
            p.add_run('📞 ').bold = True
            p.add_run(contact)

        doc.add_paragraph()

        # Summary section
        if summary:
            h = doc.add_paragraph()
            hr = h.add_run('Résumé')
            hr.bold = True
            hr.font.color.rgb = RGBColor(23,162,255)
            doc.add_paragraph(summary)

        # Experiences
        if experiences:
            h = doc.add_paragraph()
            hh = h.add_run('Expériences')
            hh.bold = True
            hh.font.color.rgb = RGBColor(23,162,255)
            for exp in experiences:
                role = exp.get('role','')
                comp = exp.get('company','')
                period = f"{exp.get('from','')} - {exp.get('to','')}" if exp.get('from') or exp.get('to') else ''
                p = doc.add_paragraph(style='List Number')
                r = p.add_run(f"{role} — {comp} {('(' + period + ')') if period else ''}")
                r.bold = True
                desc = exp.get('description','')
                if desc:
                    doc.add_paragraph(desc, style='List Bullet')

        # Education
        if education:
            h = doc.add_paragraph()
            hh = h.add_run('Éducation')
            hh.bold = True
            hh.font.color.rgb = RGBColor(23,162,255)
            for edu in education:
                deg = edu.get('degree','')
                inst = edu.get('institution','')
                year = edu.get('year','')
                p = doc.add_paragraph()
                rr = p.add_run(f"• {deg} — {inst} {('(' + str(year) + ')') if year else ''}")

        # Skills as chips (bullets)
        if skills:
            h = doc.add_paragraph()
            hh = h.add_run('Compétences')
            hh.bold = True
            hh.font.color.rgb = RGBColor(23,162,255)
            p = doc.add_paragraph()
            p.add_run(', '.join(skills))

        # Footer timestamp
        doc.add_paragraph()
        doc.add_paragraph(f'Généré: {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}')

        # Save DOCX
        fname = f"cv_{secure_filename(name)}_{int(datetime.utcnow().timestamp()*1000)}.docx"
        out_path = os.path.join(UPLOAD_DIR, fname)
        doc.save(out_path)

        result = {'generated': True, 'docx': fname, 'download_docx': f"/modules/saas/download/{fname}"}

        # Try to also export to PDF via LibreOffice if available
        try:
            ok = _libreoffice_convert(out_path, UPLOAD_DIR, 'pdf')
            pdf_name = f"{os.path.splitext(fname)[0]}.pdf"
            pdf_path = os.path.join(UPLOAD_DIR, pdf_name)
            if ok and os.path.exists(pdf_path):
                result['pdf'] = pdf_name
                result['download_pdf'] = f"/modules/saas/download/{pdf_name}"
        except Exception:
            pass

        Conversion.create(original=name, output=fname, direction='cv-professional', status='done')
        return jsonify(result)
    except Exception:
        current_app.logger.exception('CV generation failed')
        return jsonify({'error': 'cv generation failed'}), 500


@saas_bp.route('/calculator/calc', methods=['POST'])
@premium_required
def calculator():
    data = request.json or {}
    calc_type = (data.get('type') or '').lower()
    if calc_type == 'loan':
        try:
            principal = float(data.get('principal', 0))
            annual_rate = float(data.get('annual_rate', 0))
            years = float(data.get('years', 0))
            if principal <= 0 or years <= 0:
                return jsonify({'error': 'invalid principal or years'}), 400
            monthly_rate = annual_rate / 100.0 / 12.0
            n = years * 12
            if monthly_rate == 0:
                payment = principal / n
            else:
                payment = principal * (monthly_rate * (1+monthly_rate)**n) / ((1+monthly_rate)**n - 1)
            return jsonify({'type': 'loan', 'monthly_payment': round(payment, 2), 'periods': int(n)})
        except Exception:
            return jsonify({'error': 'calculation failed'}), 400
    elif calc_type == 'bmi':
        try:
            weight = float(data.get('weight_kg', 0))
            height_cm = float(data.get('height_cm', 0))
            if weight <= 0 or height_cm <= 0:
                return jsonify({'error': 'invalid weight or height'}), 400
            h = height_cm / 100.0
            bmi = weight / (h*h)
            return jsonify({'type': 'bmi', 'bmi': round(bmi, 2)})
        except Exception:
            return jsonify({'error': 'calculation failed'}), 400
    else:
        return jsonify({'error': 'unsupported calculator type'}), 400


# ============ CONVERTISSEUR DE DEVISE ============
# Force reload - v2

@saas_bp.route('/currency-converter', methods=['GET'])
def currency_converter():
    """Affiche la page du convertisseur de devise"""
    try:
        return render_template('currency_converter.html', _current_module='saas')
    except Exception as e:
        return f"Error: {str(e)}", 500


@saas_bp.route('/api/currency/convert', methods=['POST'])
def convert_currency():
    """API pour convertir une devise à une autre"""
    data = request.get_json() or {}
    
    amount = float(data.get('amount', 0))
    from_curr = data.get('from', 'EUR').upper()
    to_curr = data.get('to', 'USD').upper()
    
    if amount < 0:
        return jsonify({'error': 'amount cannot be negative'}), 400
    
    # Taux de change simulés (à remplacer par une API réelle)
    exchange_rates = {
        'EUR': 1.0,
        'USD': 1.12,
        'GBP': 0.86,
        'JPY': 150.5,
        'CHF': 0.98,
        'CAD': 1.48,
        'AUD': 1.68,
        'CNY': 8.2,
        'INR': 93.5,
        'BRL': 5.6,
        'XOF': 655.957,
        'MXN': 19.2,
        'SGD': 1.52,
        'HKD': 8.75,
        'NOK': 11.5,
        'SEK': 11.8,
        'DKK': 7.46,
        'NZD': 1.85,
    }
    
    from_rate = exchange_rates.get(from_curr)
    to_rate = exchange_rates.get(to_curr)
    
    if not from_rate or not to_rate:
        return jsonify({
            'error': 'unsupported currency',
            'supported': list(exchange_rates.keys())
        }), 400
    
    result = (amount / from_rate) * to_rate
    rate = to_rate / from_rate
    
    return jsonify({
        'success': True,
        'amount': amount,
        'from_currency': from_curr,
        'to_currency': to_curr,
        'result': round(result, 2),
        'exchange_rate': round(rate, 4),
        'timestamp': datetime.now().isoformat()
    })


@saas_bp.route('/api/currency/rates', methods=['GET'])
def get_exchange_rates():
    """Retourne tous les taux de change disponibles"""
    exchange_rates = {
        'EUR': 1.0,
        'USD': 1.12,
        'GBP': 0.86,
        'JPY': 150.5,
        'CHF': 0.98,
        'CAD': 1.48,
        'AUD': 1.68,
        'CNY': 8.2,
        'INR': 93.5,
        'BRL': 5.6,
        'XOF': 655.957,
        'MXN': 19.2,
        'SGD': 1.52,
        'HKD': 8.75,
        'NOK': 11.5,
        'SEK': 11.8,
        'DKK': 7.46,
        'NZD': 1.85,
    }
    
    return jsonify({
        'success': True,
        'base_currency': 'EUR',
        'rates': exchange_rates,
        'timestamp': datetime.now().isoformat()
    })
